"""Parses uploaded files (CSV, JSON, Excel, DOCX, SQLite) into text for LLM analysis."""

import json
import csv
import io
import logging
import sqlite3
import tempfile
import os

logger = logging.getLogger(__name__)


def parse_csv(content: bytes) -> str:
    text = content.decode("utf-8")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return "Empty CSV file."

    header = rows[0]
    data_rows = rows[1:]
    preview = data_rows[:50]

    lines = [
        f"CSV file with {len(data_rows)} rows and {len(header)} columns.",
        f"Columns: {', '.join(header)}",
        "",
        "Data:",
    ]
    lines.append(" | ".join(header))
    lines.append(" | ".join(["---"] * len(header)))
    for row in preview:
        lines.append(" | ".join(row))

    if len(data_rows) > 50:
        lines.append(f"\n... ({len(data_rows) - 50} more rows not shown)")

    return "\n".join(lines)


def parse_json(content: bytes) -> str:
    data = json.loads(content.decode("utf-8"))

    if isinstance(data, list):
        preview = data[:50]
        text = f"JSON array with {len(data)} items.\n\nFirst {len(preview)} items:\n"
        text += json.dumps(preview, indent=2, default=str)
        if len(data) > 50:
            text += f"\n\n... ({len(data) - 50} more items not shown)"
        return text
    elif isinstance(data, dict):
        text = "JSON object with keys: " + ", ".join(data.keys()) + "\n\n"
        text += json.dumps(data, indent=2, default=str)
        return text
    else:
        return str(data)


def parse_excel(content: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    lines = [f"Excel file with {len(wb.sheetnames)} sheet(s): {', '.join(wb.sheetnames)}"]

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            lines.append(f"\n## Sheet: {sheet_name}\nEmpty sheet.")
            continue

        header = [str(c) if c is not None else "" for c in rows[0]]
        data_rows = rows[1:]
        preview = data_rows[:50]

        lines.append(f"\n## Sheet: {sheet_name} ({len(data_rows)} rows, {len(header)} columns)")
        lines.append(" | ".join(header))
        lines.append(" | ".join(["---"] * len(header)))
        for row in preview:
            lines.append(" | ".join(str(c) if c is not None else "" for c in row))

        if len(data_rows) > 50:
            lines.append(f"\n... ({len(data_rows) - 50} more rows not shown)")

    wb.close()
    return "\n".join(lines)


def parse_docx(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    lines = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            level = para.style.name[-1] if para.style.name[-1].isdigit() else "1"
            lines.append(f"{'#' * int(level)} {para.text}")
        elif para.style.name == "List Bullet":
            lines.append(f"- {para.text}")
        elif para.text.strip():
            lines.append(para.text)

    for i, table in enumerate(doc.tables):
        lines.append(f"\n**Table {i + 1}:**")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            lines.append(" | ".join(cells))
            if row_idx == 0:
                lines.append(" | ".join(["---"] * len(cells)))

    return "\n".join(lines)


def parse_sqlite(content: bytes) -> str:
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".db")
    try:
        os.write(tmp_fd, content)
        os.close(tmp_fd)

        conn = sqlite3.connect(tmp_path)
        cursor = conn.cursor()

        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
        tables = [row[0] for row in cursor.fetchall()]

        if not tables:
            conn.close()
            return "Empty SQLite database (no tables found)."

        lines = [f"SQLite database with {len(tables)} table(s): {', '.join(tables)}"]

        for table in tables:
            cursor.execute(f"PRAGMA table_info('{table}')")
            columns = [row[1] for row in cursor.fetchall()]

            cursor.execute(f"SELECT COUNT(*) FROM '{table}'")
            row_count = cursor.fetchone()[0]

            cursor.execute(f"SELECT * FROM '{table}' LIMIT 50")
            rows = cursor.fetchall()

            lines.append(f"\n## Table: {table} ({row_count} rows, {len(columns)} columns)")
            lines.append(" | ".join(columns))
            lines.append(" | ".join(["---"] * len(columns)))
            for row in rows:
                lines.append(" | ".join(str(v) if v is not None else "" for v in row))

            if row_count > 50:
                lines.append(f"\n... ({row_count - 50} more rows not shown)")

        conn.close()
        return "\n".join(lines)
    finally:
        os.unlink(tmp_path)


PARSERS = {
    ".csv": parse_csv,
    ".json": parse_json,
    ".xlsx": parse_excel,
    ".xls": parse_excel,
    ".docx": parse_docx,
    ".db": parse_sqlite,
}

SUPPORTED_EXTENSIONS = set(PARSERS.keys())


def parse_file(filename: str, content: bytes) -> str:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parser = PARSERS.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    return parser(content)
